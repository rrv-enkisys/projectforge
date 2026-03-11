from __future__ import annotations

"""SOW text extraction and section detection."""
import io
import logging
import re
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# Patterns that indicate the start of a SOW section heading
_SECTION_PATTERNS: dict[str, re.Pattern[str]] = {
    "scope": re.compile(
        r"scope(?:\s+of\s+work)?|project\s+scope",
        re.IGNORECASE,
    ),
    "deliverables": re.compile(
        r"deliverables?|work\s+products?|outputs?",
        re.IGNORECASE,
    ),
    "timeline": re.compile(
        r"timeline|schedule|project\s+schedule|milestones?|duration|phases?",
        re.IGNORECASE,
    ),
    "budget": re.compile(
        r"budget|cost|pricing|payment|compensation|fees?|financial",
        re.IGNORECASE,
    ),
    "terms": re.compile(
        r"terms?\s+(?:and\s+conditions?|of\s+service)|legal|liability"
        r"|ip\s+rights?|confidentiality|warranty",
        re.IGNORECASE,
    ),
}

# Maximum line length considered a section heading (long lines are body text)
_MAX_HEADING_LEN = 120


@dataclass
class ExtractedDocument:
    """Holds the full plain text and detected sections of a SOW."""

    full_text: str
    sections: dict[str, str] = field(default_factory=dict)


class SOWExtractor:
    """Extracts text from various file formats and detects SOW sections."""

    async def extract(self, file_content: bytes, content_type: str) -> ExtractedDocument:
        """Extract text from *file_content* and detect SOW sections.

        Args:
            file_content: Raw file bytes.
            content_type: MIME type or file extension (e.g. ``"application/pdf"``).

        Returns:
            :class:`ExtractedDocument` with full text and section map.
        """
        text = await self._extract_text(file_content, content_type)
        sections = self._detect_sections(text)
        return ExtractedDocument(full_text=text, sections=sections)

    # ------------------------------------------------------------------
    # Text extraction helpers
    # ------------------------------------------------------------------

    async def _extract_text(self, content: bytes, content_type: str) -> str:
        ct = content_type.lower()

        if ct in ("text/plain", ".txt"):
            return content.decode("utf-8", errors="replace")

        if ct in ("text/markdown", ".md"):
            return content.decode("utf-8", errors="replace")

        if ct in ("application/pdf", ".pdf"):
            return await self._extract_pdf(content)

        if ct in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".docx",
        ):
            return await self._extract_docx(content)

        # Fallback: try UTF-8 decode
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            logger.warning("Could not decode content with type=%s", content_type)
            return ""

    @staticmethod
    async def _extract_pdf(content: bytes) -> str:
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(content)) as pdf:
                parts = [p.extract_text() for p in pdf.pages if p.extract_text()]
            return "\n\n".join(parts)
        except ImportError:
            pass

        try:
            import PyPDF2

            reader = PyPDF2.PdfReader(io.BytesIO(content))
            parts = [page.extract_text() for page in reader.pages if page.extract_text()]
            return "\n\n".join(parts)
        except ImportError:
            logger.error("Neither pdfplumber nor PyPDF2 installed")
            raise ImportError("PDF extraction requires pdfplumber or PyPDF2")

    @staticmethod
    async def _extract_docx(content: bytes) -> str:
        try:
            import docx

            doc = docx.Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(parts)
        except ImportError:
            logger.error("python-docx not installed")
            raise ImportError("DOCX extraction requires python-docx")

    # ------------------------------------------------------------------
    # Section detection
    # ------------------------------------------------------------------

    @staticmethod
    def _detect_sections(text: str) -> dict[str, str]:
        """Split text into labelled sections using heading pattern matching.

        A line is treated as a section heading when:
        - It is short (< 120 chars), AND
        - The matched keyword covers at least 40 % of the line length.

        This prevents body-text lines that happen to contain a keyword
        (e.g. "Total project cost: $45,000 USD") from being treated as headings.
        """
        sections: dict[str, str] = {}
        current_section: str | None = None
        current_lines: list[str] = []

        for line in text.splitlines():
            stripped = line.strip()
            matched: str | None = None

            if stripped and len(stripped) < _MAX_HEADING_LEN:
                for section_name, pattern in _SECTION_PATTERNS.items():
                    m = pattern.search(stripped)
                    if m and len(m.group(0)) >= len(stripped) * 0.4:
                        matched = section_name
                        break

            if matched:
                # Flush the previous section buffer
                if current_section and current_lines:
                    prev = sections.get(current_section, "")
                    sections[current_section] = (
                        (prev + "\n" + "\n".join(current_lines)).strip()
                    )
                current_section = matched
                current_lines = []
            else:
                current_lines.append(line)

        # Flush the last section
        if current_section and current_lines:
            prev = sections.get(current_section, "")
            sections[current_section] = (
                (prev + "\n" + "\n".join(current_lines)).strip()
            )

        return sections


# ---------------------------------------------------------------------------
# Singleton
# ---------------------------------------------------------------------------
_sow_extractor: SOWExtractor | None = None


def get_sow_extractor() -> SOWExtractor:
    """Return the shared :class:`SOWExtractor` singleton."""
    global _sow_extractor
    if _sow_extractor is None:
        _sow_extractor = SOWExtractor()
    return _sow_extractor
